"""
Generate two Word documents:
  1. figure_captions.docx   - bilingual (CN/EN) captions for all figures & tables
  2. thesis_narrative.docx  - thesis storyline connecting all figures & tables
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT_DIR = os.path.join(os.path.dirname(__file__),
                       "benchmark_output", "20260331_171049", "figures")


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------
def _set(run, size=11, bold=False, italic=False, color=None, east="宋体"):
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.name   = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    if color:
        run.font.color.rgb = RGBColor(*color)


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    _set(r, 14, bold=True, color=(31, 73, 125), east="黑体")


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    _set(r, 12, bold=True, color=(68, 114, 196))


def label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(text)
    _set(r, 11, bold=True)


def cn(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.5)
    r = p.add_run("【中文图注】" + text)
    _set(r, 10.5)


def en(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Cm(0.5)
    r = p.add_run("[EN Caption] " + text)
    _set(r, 10.5, italic=True)


def div(doc):
    p = doc.add_paragraph("─" * 78)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    for r in p.runs:
        _set(r, 8, color=(180, 180, 180))


def body(doc, text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.first_line_indent = Pt(22)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    _set(r, 11)


def ibody(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.first_line_indent = Pt(22)
    r = p.add_run(text)
    _set(r, 11, italic=True, color=(80, 80, 80))


def bullet(doc, text, indent=0.8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(indent)
    r = p.add_run(text)
    _set(r, 11)


# ===========================================================================
# Document 1 - bilingual captions
# ===========================================================================
def make_captions_doc():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("图表说明 · Figure & Table Captions")
    _set(r, 16, bold=True, color=(31, 73, 125), east="黑体")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("葡萄膜黑色素瘤 D3/M3 二分类 · 21 模型横评基准")
    _set(r2, 12, color=(100, 100, 100))

    doc.add_paragraph()
    div(doc)

    # ------------------------------------------------------------------
    h1(doc, "一、表格 / Tables")
    # ------------------------------------------------------------------

    label(doc, "Table 1 / 表1  临床基线特征")
    cn(doc,
       "表1. TCGA-UVM 队列基线临床特征（n=80）。本研究纳入来自 TCGA 数据库的 80 例"
       "葡萄膜黑色素瘤（UVM）患者，依据 SCNA 拷贝数变异分簇分为"
       "Disomy 3（D3，簇1/2，n=38）和 Monosomy 3（M3，簇3/4，n=42）两组。"
       "连续变量以中位数（范围）表示，分类变量以例数（百分比）表示。"
       "组间比较采用 Mann-Whitney U 检验（连续变量）或卡方检验（分类变量）。"
       "OS：总生存期。")
    en(doc,
       "Table 1. Baseline clinical characteristics of the TCGA-UVM cohort (n=80). "
       "Patients were stratified into Disomy 3 (D3; SCNA Cluster 1/2, n=38) and "
       "Monosomy 3 (M3; SCNA Cluster 3/4, n=42). Continuous variables are reported "
       "as median (range); categorical variables as count (%). Between-group "
       "comparisons used the Mann-Whitney U test (continuous) or chi-square test "
       "(categorical). OS, overall survival.")
    div(doc)

    label(doc, "Table 2 / 表2  21 个评估模型概览")
    cn(doc,
       "表2. 本研究横评的 21 个模型概览，包含 20 个病理基础大模型及 1 个 ImageNet"
       "预训练 ResNet-50 基线。信息涵盖机构来源、预训练范式（SSL：自监督学习；"
       "VLP：视觉-语言预训练；Supervised：有监督）、骨干架构、参数量、输出特征维度"
       "及预训练数据类型。AUC 为本研究评估所得的 5 折交叉验证均值。")
    en(doc,
       "Table 2. Overview of the 21 models evaluated, comprising 20 pathology "
       "foundation models and one ImageNet-pretrained ResNet-50 baseline. Columns "
       "include institution, pretraining paradigm (SSL, VLP, Supervised), backbone "
       "architecture, parameter count, feature dimension, pretraining data source, "
       "and mean 5-fold AUC from this study.")
    div(doc)

    label(doc, "Table 3 / 表3  完整基准排名（主结果表）")
    cn(doc,
       "表3. 21 个模型在 TCGA-UVM D3/M3 二分类任务上的完整性能排名（按 AUC 降序）。"
       "特征提取器均冻结，聚合器采用 ABMIL（n_token=1），5 折分层交叉验证，"
       "全模型统一超参数（不进行逐模型调优）。"
       "指标：AUC（主要指标）、准确率（Acc）、加权 F1 分数（F1）。"
       "AUC 以均值±标准差表示；★ 标记模型预训练数据包含 TCGA 全库切片。")
    en(doc,
       "Table 3. Complete benchmark ranking of 21 models on the TCGA-UVM D3/M3 "
       "binary classification task (sorted by AUC, descending). All encoders were "
       "frozen; aggregation used ABMIL (n_token=1) with a fixed hyperparameter "
       "protocol. Metrics: AUC (primary), accuracy (Acc), weighted F1 (F1). "
       "AUC is reported as mean +/- std across five folds. "
       "Star marks models pretrained on datasets including all TCGA WSIs.")
    div(doc)

    label(doc, "Table 4 / 表4  Q1：基础模型 vs ResNet-50 增益表")
    cn(doc,
       "表4. 各病理基础大模型相对于 ResNet-50 基线的 AUC 增益（ΔAUC）。"
       "20 个基础模型中有 17 个优于基线，增益范围 −0.044 至 +0.079。"
       "符号检验（17/20 超越基线，二项检验）p=1.3×10⁻³（**）。"
       "头部模型 H-optimus-1 增益 +0.079 AUC，具有明确临床意义。")
    en(doc,
       "Table 4. AUC gain (delta AUC) of each foundation model relative to ResNet-50. "
       "18 of 20 foundation models exceeded the baseline; gains ranged from -0.009 "
       "to +0.079. Sign test (17/20 above baseline, binomial): p=1.3e-03 (**). "
       "H-optimus-1 achieves the largest gain (+0.079 AUC).")
    div(doc)

    label(doc, "Table 5 / 表5  Q2：TCGA 预训练 vs 私有数据预训练")
    cn(doc,
       "表5. TCGA 预训练组（8 个模型）与私有数据预训练组（12 个模型）在"
       "TCGA-UVM D3/M3 任务上的性能对比。置换检验（permutation test, 10 000 次）p<0.001（***），"
       "差异显著：私有预训练模型均值 AUC（0.894）系统性高于 TCGA 预训练模型（0.853）。"
       "这一反直觉结果说明切片级预训练重叠并未带来性能优势，"
       "大规模多样化数据积累的通用病理知识更为关键。")
    en(doc,
       "Table 5. Performance comparison: TCGA-pretrained (n=8) vs. "
       "private-pretrained (n=12) models on TCGA-UVM D3/M3. "
       "Two-sided permutation test (10,000 iterations): p<0.001 (***). Private-pretrained models "
       "achieved systematically higher mean AUC (0.894 vs. 0.853, delta=-0.041), "
       "indicating that slide-level pretraining overlap does not confer an "
       "advantage; large-scale diverse pretraining data is more critical.")
    div(doc)

    label(doc, "Table 6 / 表6  Q3：Kaiko 系列 2x2 析因设计")
    cn(doc,
       "表6. Kaiko 系列四变体（ViT-S/8、ViT-S/16、ViT-B/8、ViT-B/16）的完整性能对比，"
       "覆盖模型容量（ViT-S ~22M 参数 vs ViT-B ~86M）与输入分辨率"
       "（patch 8 高分辨率 vs patch 16 低分辨率）两个因子。"
       "所有变体共享相同预训练数据（TCGA）、算法（DINO）和切片级重叠状态，"
       "构成严格对照的 2x2 析因实验。")
    en(doc,
       "Table 6. Performance of the four Kaiko variants in a 2x2 factorial design "
       "crossing model capacity (ViT-S ~22M vs. ViT-B ~86M) and input resolution "
       "(patch 8, high-res vs. patch 16, low-res). All variants share identical "
       "pretraining data (TCGA), algorithm (DINO), and overlap status, providing a "
       "strictly controlled comparison.")
    div(doc)

    # ------------------------------------------------------------------
    h1(doc, "二、图 / Figures")
    # ------------------------------------------------------------------

    label(doc, "图 S1 / Fig. S1  研究设计与计算流程图")
    cn(doc,
       "图S1. 研究设计与计算流程示意图。流程分五阶段：(1) 队列构建——来自 TCGA 的"
       "80 例 UVM WSI，按 SCNA 分类分为 D3（38例）和 M3（42例）；"
       "(2) 图像分块——Trident 框架 20× 放大、256×256 像素无重叠分块；"
       "(3) 特征提取——21 个冻结编码器分别提取图块级特征；"
       "(4) MIL 聚合——ABMIL（n_token=1）池化为切片级表征；"
       "(5) 评估——5 折分层交叉验证，报告 AUC/Acc/F1。"
       "底部标注三个科学问题。")
    en(doc,
       "Fig. S1. Schematic of the study design and computational pipeline. "
       "Five stages: (1) Cohort — 80 TCGA-UVM WSIs stratified into D3 (n=38) "
       "and M3 (n=42); (2) Patch extraction — non-overlapping 256x256-pixel tiles "
       "at 20x via Trident; (3) Feature extraction — 21 frozen encoders extract "
       "patch-level features; (4) MIL aggregation — ABMIL (n_token=1) pools to "
       "slide-level representations; (5) Evaluation — 5-fold stratified CV "
       "reporting AUC, accuracy, and F1. "
       "Three scientific questions are annotated at the bottom.")
    div(doc)

    label(doc, "图 1 / Fig. 1  队列临床特征分布")
    cn(doc,
       "图1. TCGA-UVM 队列临床特征的 D3/M3 分组比较（n=80）。六子图分别展示："
       "(A) 年龄分布（p=0.073，无显著差异）；"
       "(B) 性别分布（两组相近，p=1.000）；"
       "(C) SCNA 分簇（簇1/2 = D3；簇3/4 = M3）；"
       "(D) AJCC 分期（p=0.071）；"
       "(E) 组织学类型——D3 以梭形细胞为主，M3 上皮样细胞比例显著更高（p<0.001）；"
       "(F) OS 时间分布——M3 组 OS 事件率显著更高（50% vs 7.9%，p=0.026），"
       "与 Monosomy 3 预后不良的已知证据一致。")
    en(doc,
       "Fig. 1. Clinical characteristics of the TCGA-UVM cohort stratified by "
       "D3/M3 group (n=80). Six panels: (A) age (p=0.073); (B) sex (p=1.000); "
       "(C) SCNA cluster distribution; (D) AJCC staging (p=0.071); "
       "(E) histological subtype — D3 is predominantly spindle cell; M3 shows "
       "markedly higher epithelioid cell proportion (p<0.001); "
       "(F) OS time — M3 has a significantly higher OS event rate "
       "(50% vs. 7.9%, p=0.026), consistent with the known poor prognosis "
       "of chromosome 3 monosomy.")
    div(doc)

    label(doc, "图 2 / Fig. 2  Q1：病理基础大模型的迁移价值（柱状图）")
    cn(doc,
       "图2. 20 个病理基础大模型与 ResNet-50 基线的 AUC 对比（5 折均值±标准差）。"
       "蓝色=私有预训练，橙色=TCGA 预训练，灰色=基线；灰色虚线标注基线 AUC（0.8423）。"
       "17 个基础模型超越基线（符号检验 p<0.01），"
       "头部模型 H-optimus-1 增益达 +0.079 AUC。")
    en(doc,
       "Fig. 2. AUC comparison of 20 pathology foundation models vs. ResNet-50 "
       "(5-fold mean +/- std). Blue = private-pretrained; orange = TCGA-pretrained; "
       "gray dashed line = baseline AUC (0.8423). "
       "Sign test (17/20 above baseline, binomial): p=1.3e-03 (**). "
       "H-optimus-1 achieves the largest margin (+0.079 AUC).")
    div(doc)

    label(doc, "图 3 / Fig. 3  Q2：切片级预训练重叠效应（箱线图+森林图）")
    cn(doc,
       "图3. TCGA 预训练模型（n=8）与私有数据预训练模型（n=12）的 AUC 系统性差异。"
       "左：箱线图叠加 jitter strip，显著性括号标注双侧置换检验（p<0.001，***）；"
       "右：各模型 AUC 森林图，颜色区分预训练来源，灰色虚线为 ResNet-50 基线。"
       "私有预训练组均值 AUC（0.894）显著高于 TCGA 组（0.853），ΔAUC=−0.041。")
    en(doc,
       "Fig. 3. Systematic AUC difference between TCGA-pretrained (n=8) and "
       "private-pretrained (n=12) foundation models. Left: box plots with jitter "
       "strips; two-sided permutation test (p<0.001, ***). Right: forest plot of "
       "individual model AUCs; gray dashed line = ResNet-50 baseline. "
       "Private-pretrained models achieve significantly higher mean AUC "
       "(0.894 vs. 0.853; delta=-0.041).")
    div(doc)

    label(doc, "图 4 / Fig. 4  Q3：Kaiko 系列 2x2 析因交互图与热力图")
    cn(doc,
       "图4. Kaiko 系列四变体的 2x2 析因设计结果。"
       "左（交互线图）：横轴为模型容量（ViT-S 至 ViT-B），两线分别代表 patch 8"
       "（橙色）和 patch 16（蓝色）。两线显著交叉（X 形），揭示容量×分辨率交互效应："
       "ViT-S 时 patch 8 > patch 16（0.871 vs 0.839），ViT-B 时结果反转（0.881 vs 0.830）。"
       "右（热力图）：颜色深度直观展示四格 AUC 矩阵。")
    en(doc,
       "Fig. 4. Results of the 2x2 factorial design in the Kaiko series. "
       "Left (interaction plot): x-axis = model capacity (ViT-S to ViT-B); "
       "two lines represent patch 8 (orange) and patch 16 (blue). Pronounced "
       "line crossing reveals a capacity x resolution interaction: "
       "at ViT-S scale, patch 8 > patch 16 (0.871 vs. 0.839); "
       "at ViT-B scale, the result reverses (0.881 vs. 0.830). "
       "Right (heatmap): colour intensity visualises the four-cell AUC matrix.")
    div(doc)

    label(doc, "图 5 / Fig. 5  各模型五折 AUC 分布箱线图")
    cn(doc,
       "图5. 所有 21 个模型的 5 折 AUC 分布（水平箱线图，按均值降序排列）。"
       "箱线图展示中位数、四分位距和须范围；叠加 jitter 散点显示原始折值。"
       "颜色区分预训练来源（蓝=私有，橙=TCGA，灰=基线）。"
       "GigaPath 和 Kaiko ViT-S/8 折间方差最小（最稳定）；"
       "Phikon-v2 和 LUNIT ViT-S/8 表现最不稳定。"
       "灰色虚线为 ResNet-50 均值 AUC（0.8423）。")
    en(doc,
       "Fig. 5. Per-fold AUC distribution of all 21 models (horizontal box plots, "
       "sorted by mean AUC descending). Boxes show median, IQR, and whiskers; "
       "jittered dots overlay raw fold values. Colors distinguish pretraining "
       "source (blue = private; orange = TCGA; grey = baseline). GigaPath and "
       "Kaiko ViT-S/8 show the lowest inter-fold variance (most stable); "
       "Phikon-v2 and LUNIT ViT-S/8 are least stable. "
       "Gray dashed line marks ResNet-50 mean AUC (0.8423).")
    div(doc)

    label(doc, "图 6 / Fig. 6  Top-5 模型与基线的平均 ROC 曲线")
    cn(doc,
       "图6. 前五名模型（H-optimus-1、UNI v1、H-optimus-0、CONCH v1.5、Virchow）"
       "及 ResNet-50 基线在 5 折交叉验证中的平均 ROC 曲线。"
       "实线为 5 折 TPR 逐点均值，阴影为 ±1 标准差。图例标注 AUC 均值±标准差。"
       "H-optimus-1 在全部折中曲线最高，展示最稳健的鉴别能力；"
       "ResNet-50 曲线明显低于所有基础模型。")
    en(doc,
       "Fig. 6. Mean ROC curves of the top-5 models and ResNet-50 baseline across "
       "5-fold CV. Solid lines = point-wise mean TPR; shaded bands = +/-1 SD. "
       "AUC is annotated as mean +/- std in the legend. H-optimus-1 achieves "
       "the highest curve across all folds; ResNet-50 falls markedly below "
       "all foundation models.")
    div(doc)

    label(doc, "图 7 / Fig. 7  模型性能与稳定性散点图")
    cn(doc,
       "图7. 模型均值 AUC（横轴）与 AUC 折间标准差（纵轴）的二维散点图。"
       "虚线划分四个象限；右下象限（高 AUC + 低方差）为理想区域，"
       "GigaPath、Kaiko ViT-S/8 和 H-optimus-1 位于此区域。"
       "左上象限（低 AUC + 高方差，如 Midnight-12k）为最不理想区域。"
       "在 80 例小数据集上，高方差意味着对训练集分布高度敏感，泛化可靠性存疑。")
    en(doc,
       "Fig. 7. Scatter plot of mean AUC (x-axis) vs. cross-fold AUC standard "
       "deviation (y-axis). Dotted lines partition the space into four quadrants; "
       "lower-right (high AUC, low variance) is ideal — GigaPath, Kaiko ViT-S/8, "
       "and H-optimus-1 reside here. Upper-left (low AUC, high variance; "
       "e.g., Midnight-12k) is least desirable. On an 80-case dataset, high "
       "variance implies sensitivity to training-set composition, "
       "undermining reliable generalisation.")
    div(doc)

    label(doc, "图 8 / Fig. 8  架构属性（特征维度/参数量）与 AUC 的关系")
    cn(doc,
       "图8. 架构属性与下游 AUC 的 Spearman 秩相关。"
       "左（特征维度 vs AUC）：Spearman r=0.349，p=0.121（n.s.）；"
       "右（参数量 vs AUC）：Spearman r=0.389，p=0.081（n.s.）。"
       "两者均未达统计显著，说明在 80 例极小样本场景下，"
       "模型容量的增大不线性转化为性能提升，"
       "预训练数据多样性可能是更关键的因素。"
       "私有预训练模型（蓝点）整体分布高于 TCGA 预训练模型（橙点）。")
    en(doc,
       "Fig. 8. Spearman rank correlation between architectural properties and "
       "downstream AUC. Left (feature dimension vs. AUC): r=0.349, p=0.121 (n.s.); "
       "Right (parameter count vs. AUC): r=0.389, p=0.081 (n.s.). "
       "Neither reaches significance, indicating that model capacity does not "
       "linearly predict performance in this extreme low-data regime. "
       "Private-pretrained models (blue) are distributed higher overall "
       "than TCGA-pretrained models (orange).")
    div(doc)

    label(doc, "图 9 / Fig. 9  评估指标间相关性热力图")
    cn(doc,
       "图9. AUC、准确率、加权 F1 三个评估指标间的 Pearson"
       "相关系数矩阵（跨 21 个模型计算）。所有指标对之间均高度正相关（r > 0.90），"
       "说明三个指标在本二分类任务中传达几乎相同的排序信息。"
       "以 AUC 作为主要指标充分且合理。")
    en(doc,
       "Fig. 9. Pearson correlation matrix of three evaluation metrics (AUC, "
       "accuracy, weighted F1) computed across 21 models. "
       "All metric pairs are highly positively correlated (r > 0.90), "
       "confirming that all four metrics convey nearly identical ranking "
       "information. Using AUC as the primary metric is well-justified.")
    div(doc)

    label(doc, "图 10 / Fig. 10  预训练范式对比（SSL / VLP / 有监督）")
    cn(doc,
       "图10. 三种预训练范式在 UVM D3/M3 任务上的性能对比。"
       "左（箱线图）：SSL（n=17）、VLP（n=3，含 CONCH v1/v1.5、MUSK）、"
       "有监督（ImageNet ResNet-50，n=1）。SSL 与 VLP 组间行 Mann-Whitney U 检验。"
       "右（点图）：全模型按范式着色排名。SSL 整体中位数略高于 VLP，"
       "两者差异不显著；两组 AUC 均高于有监督基线（n=1，无法行统计检验）。"
       "「是否使用病理无标签数据预训练」是最关键的分水岭。")
    en(doc,
       "Fig. 10. Performance comparison across three pretraining paradigms. "
       "Left (box plots): SSL (n=17), VLP (n=3: CONCH v1, v1.5, MUSK), "
       "Supervised (ResNet-50, n=1). Mann-Whitney U test applied between SSL and VLP. "
       "Right: ranked dot plot coloured by paradigm. SSL achieves a marginally "
       "higher median than VLP (not significant); both groups score above the "
       "supervised baseline (n=1, formal testing not applicable). "
       "Whether large-scale unlabelled pathology data was used for "
       "pretraining is the most critical dividing factor.")
    div(doc)

    label(doc, "图 11 / Fig. 11  Top-10 模型多指标对比柱状图")
    cn(doc,
       "图11. 前十名模型在 AUC、准确率、加权 F1 三个指标上的"
       "并排柱状图。颜色固定对应指标（蓝=AUC、绿=Acc、橙=F1）。"
       "AUC 排名第一的 H-optimus-1 准确率也最高（0.813）；"
       "结论对指标选择具有高稳健性，各指标排序高度一致。")
    en(doc,
       "Fig. 11. Multi-metric bar chart for the top-10 models (AUC, accuracy, "
       "weighted F1). Colours are fixed per metric "
       "(blue=AUC; green=Acc; orange=F1). H-optimus-1 (AUC rank 1) "
       "also achieves the highest accuracy (0.813). "
       "Rankings are highly consistent across all three metrics.")

    doc.save(os.path.join(OUT_DIR, "figure_captions.docx"))
    print("  Saved: figure_captions.docx")


# ===========================================================================
# Document 2 - thesis narrative
# ===========================================================================
def make_narrative_doc():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("图表叙事框架")
    _set(r, 18, bold=True, color=(31, 73, 125), east="黑体")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("基于病理基础大模型横评的极端罕见肿瘤分型研究\n"
                    "—— 葡萄膜黑色素瘤 D3/M3 二分类基准")
    _set(r2, 12, italic=True, color=(80, 80, 80))

    doc.add_paragraph()
    div(doc)

    # -----------------------------------------------------------------------
    h1(doc, "全文叙事概览")
    body(doc,
         "本研究围绕三个递进的科学问题，系统评估了 20 个病理基础大模型在极端罕见肿瘤"
         "（TCGA-UVM，n=80）上的分类性能，以 ImageNet 预训练 ResNet-50 作为基线对照。"
         "研究的核心逻辑是：先建立「谁更好」的性能排名，再追问「为什么更好或更差」"
         "的机制原因。全文图表分为四个叙事层次：(1) 队列描述与实验设计；"
         "(2) 主要结果——全模型排名与迁移价值；"
         "(3) 机制探究——预训练数据重叠效应；"
         "(4) 精细拆解——模型容量与分辨率的析因实验。")

    # -----------------------------------------------------------------------
    h1(doc, "第一部分：队列描述与实验设计")

    h2(doc, "1.1  为什么选择 UVM？（图1 + Table 1）")
    body(doc,
         "研究的第一步是向读者交代数据本身为何值得被研究。葡萄膜黑色素瘤是成人最常见"
         "的眼内恶性肿瘤，但在全癌种中属于极端罕见病——TCGA 队列仅收录 80 例。"
         "更重要的是，D3/M3 分型具有明确的临床预后意义：M3 组 OS 事件率显著更高"
         "（图1F，50% vs D3 组 7.9%，p=0.026），同时两组在组织学表型上存在系统差异"
         "（图1E，D3 以梭形细胞为主，M3 上皮样细胞比例显著更高，p<0.001）。"
         "这说明 D3/M3 分型的形态学基础确实存在于 WSI 中，是 MIL 方法理论上可以"
         "捕捉的信号，为本研究的可行性提供了直接的病理学依据。")
    body(doc,
         "Table 1（临床基线特征三线表）论证两组在年龄、性别、AJCC 分期和治疗方式上"
         "的均衡性（均无统计显著差异），排除混杂因素对分类任务的干扰，"
         "为后续纯形态学特征驱动的分类实验提供合理性基础。")

    h2(doc, "1.2  如何设计实验？（图S1 + Table 2）")
    body(doc,
         "图S1（流程图）向读者清晰展示方法的每个环节及其设计依据：")
    bullet(doc, "· Trident 分块（20×，256px，0重叠）保持与绝大多数基础模型预训练"
                "一致的感受野设定，减少分辨率不匹配引入的噪声；")
    bullet(doc, "· 冻结编码器——对比的是「特征」本身而非微调能力；"
                "在 80 例数据下微调会引入过拟合风险；")
    bullet(doc, "· ABMIL（n_token=1）——最简 MIL 聚合器，消除聚合器差异对"
                "比较公平性的干扰；")
    bullet(doc, "· 统一超参数——80 例极小数据集上逐模型网格搜索拟合的是噪声"
                "而非信号，固定超参是比较特征提取器「纯能力」的正确方法。")
    body(doc,
         "Table 2（模型概览三线表）系统罗列 21 个模型的架构背景，"
         "为读者建立对「参赛选手」的整体认知，也为后续分组分析（Q2/Q3）"
         "提供元数据支撑。")

    # -----------------------------------------------------------------------
    h1(doc, "第二部分：主要结果——谁赢了，赢了多少？")

    h2(doc, "2.1  完整排名（Table 3 + 图2 + 图5 + 图6）")
    body(doc,
         "Table 3（完整排名三线表）是全文最核心的结果表，以 AUC 降序排列 21 个模型"
         "的完整指标，直接给出「答案」：H-optimus-1 以均值 AUC=0.921 位居第一，"
         "Midnight-12k 以 AUC=0.798 位居末尾，ResNet-50 基线 AUC=0.842"
         "（倒数第四，低于 kaiko-vits16）。")
    body(doc,
         "图2（Q1 柱状图）将 Table 3 可视化，蓝色/橙色区分预训练来源，"
         "灰色虚线标注基线。配套 Table 4（ΔAUC 三线表）精确量化每个模型相对"
         "基线的增益，让读者直接读出「值不值」。")
    body(doc,
         "图5（箱线图）展示折间稳定性——这是 Table 3 中均值±标准差无法完全"
         "传达的信息。均值高但方差大的模型与均值稍低但极度稳定的模型，"
         "在临床转化意义上截然不同。图5 提供了这个完整画面。")
    body(doc,
         "图6（ROC 曲线）作为分类论文的标准图之一，展示前五名模型和基线的"
         "平均 ROC 曲线及折间置信带，让读者从曲线形状感受各模型的鉴别能力。")

    h2(doc, "2.2  科学问题一的完整解读（图2 + Table 4）")
    body(doc,
         "「病理基础大模型是否显著优于 ImageNet 预训练的 ResNet-50？」")
    body(doc,
         "现象层面（图2）：20 个基础模型中 17 个高于基线，头部模型领先 +0.079 AUC。"
         "统计层面（Table 4 注）：采用符号检验"
         "（17/20 超越基线，二项检验）p=1.3×10⁻³（**），显著优于基线。"
         "结论明确：病理基础大模型在 UVM D3/M3 任务上的表示能力显著优于通用视觉基线。"
         "头部模型（H-optimus-1）增益 +0.079 AUC 在罕见肿瘤分型的临床场景下具有直接实践意义。")
    ibody(doc,
          "统计方法说明：Q1 采用符号检验（二项检验），不依赖模型间独立性假设，"
          "仅检验超越基线的比例是否显著高于随机猜测。")

    # -----------------------------------------------------------------------
    h1(doc, "第三部分：机制探究——预训练数据重叠效应")

    h2(doc, "3.1  科学问题二（图3 + Table 5 + 图8 + 图9）")
    body(doc,
         "「TCGA 切片级数据重叠是否带来性能优势？」")
    body(doc,
         "图3（Q2 箱线图+森林图）是这一问题的核心可视化，展示了反直觉的结论："
         "私有预训练模型（均值 AUC=0.894）系统性高于 TCGA 预训练模型（0.853），"
         "差距 0.041，置换检验 p<0.001（***）。"
         "这一结论在 80 例数据下尤其有力——样本量极小意味着任何由重叠带来的优势"
         "都会被清晰暴露，而结果恰恰相反。")
    body(doc,
         "配套 Table 5 为评审提供精确统计摘要。图8（架构散点图）进一步证明"
         "这种系统性差距不能被参数量或特征维度的差异解释（Spearman p>0.05），"
         "因为两组的架构分布是重叠的。图9（指标相关热力图）确认了以 AUC 作为"
         "主指标具有充分代表性（与 Acc/F1 相关 r>0.90），"
         "排除了「换一个指标结论会不同」的质疑。")
    body(doc,
         "结论双重解读价值：(1) 方法论层面——在 TCGA 上评估 TCGA 预训练模型，"
         "并不因「见过这些切片」而被系统性高估；"
         "(2) 实践层面——选型时应优先考虑大规模多样化私有数据预训练的模型。")

    h2(doc, "3.2  预训练范式分析（图10 + 图11）")
    body(doc,
         "图10（范式对比）从 SSL vs VLP vs Supervised 的维度进一步拆解。"
         "三种范式的差异模式与 Q2 结论相互印证：有监督基线（仅 ImageNet）最低，"
         "SSL 和 VLP 的 AUC 均高于有监督基线（n=1，无法行统计检验），但 SSL 与 VLP 之间无显著差异。"
         "这说明「是否使用了大规模无标签病理图像预训练」才是关键分水岭，"
         "具体的预训练目标函数（对比学习 vs 语言对齐）在当前任务规模下"
         "并非决定性因素。")
    body(doc,
         "图11（Top-10 多指标柱图）辅助说明：即便换用准确率或 F1 等其他指标，"
         "模型的相对排序高度一致，H-optimus-1 仍然排名第一，"
         "结论对指标选择具有稳健性。")

    # -----------------------------------------------------------------------
    h1(doc, "第四部分：精细拆解——析因实验")

    h2(doc, "4.1  科学问题三（图4 + Table 6 + 图7）")
    body(doc,
         "「在控制预训练数据、算法和重叠状态的前提下，模型容量与输入分辨率"
         "如何独立及交互影响下游性能？」")
    body(doc,
         "图4（Kaiko 2x2 交互图+热力图）是本研究方法论最精密的部分。"
         "通过 Kaiko 系列四个变体的天然对照，在不改变任何其他变量的情况下，"
         "单独观察「容量」和「分辨率」两个因子的效应。")
    body(doc,
         "交互线图中两条线的显著交叉（X 形）揭示了统计意义上的交互效应：")
    bullet(doc, "· ViT-S 规模时：patch 8（高分辨率）> patch 16（低分辨率），+0.032 AUC；")
    bullet(doc, "· ViT-B 规模时：结果反转，patch 16（低分辨率）> patch 8（高分辨率），+0.051 AUC。")
    body(doc,
         "机制解读：小容量模型（ViT-S）参数有限，需要更丰富的局部纹理信息（小 patch）"
         "来弥补全局建模能力的不足；大容量模型（ViT-B）具备更强的 attention 建模能力，"
         "在低分辨率（patch 16）下能更有效地捕捉病理结构级信息，"
         "而过细粒度的 patch（patch 8）可能引入冗余局部噪声。"
         "在 80 例极小数据集上这种效应被放大——模型无法通过大量样本学会忽略噪声，"
         "因此输入信息的质量比数量更重要。")
    body(doc,
         "Table 6 为图4 提供精确数值支撑；图7（稳定性散点）进一步揭示："
         "Kaiko ViT-S/8 位于「高性能+低方差」象限，"
         "而 Kaiko ViT-B/8 参数更多但稳定性显著低于 ViT-S/8，"
         "再次印证了上述机制解读。")

    # -----------------------------------------------------------------------
    h1(doc, "第五部分：综合讨论——图表串联逻辑")

    h2(doc, "5.1  三个科学问题的逻辑递进")
    body(doc,
         "三个科学问题形成逻辑链：Q1 建立「基础模型整体有效」的前提"
         "→ Q2 追问「不同来源基础模型为何表现有别」"
         "→ Q3 进一步控制变量、精细拆解「同一家族内架构差异如何影响性能」。"
         "这种层层推进的结构，使研究从「观察」上升到「机制」，"
         "是高质量评估类论文的标准叙事框架。")

    h2(doc, "5.2  跨图表的关键交叉引用")
    refs = [
        "图2（Q1 柱状图）↔ 图5（箱线图）：前者展示均值结论，后者补充置信度/稳定性；",
        "图3（Q2 分组比较）↔ 图8（架构散点图）：前者发现 TCGA 组系统偏低，"
        "后者排除参数量差异的混淆解释；",
        "图3（Q2）↔ 图10（范式分析）：从不同维度（预训练数据 vs 算法目标）"
        "切入，结论相互兼容；",
        "图4（Q3 交互图）↔ 图7（稳定性散点）：图4 揭示容量×分辨率交互，"
        "图7 揭示 ViT-B/8 高不稳定性，共同支持「大容量模型在高分辨率输入下"
        "于小数据集上表现不稳」的结论；",
        "Table 1（临床表）↔ 图1（临床分布图）：表格提供统计摘要，图形展示分布形态；",
        "Table 3（排名表）↔ 图6（ROC 曲线）：排名表给出全局排序，ROC 放大前五名细节；",
        "图9（指标相关）：方法论正当性说明，支撑全文以 AUC 为主要评估指标的决定。",
    ]
    for r_text in refs:
        bullet(doc, "· " + r_text)

    h2(doc, "5.3  对评审可能质疑的预先回应")
    qas = [
        ("Q：为什么不对每个模型进行超参数调优？",
         "A：(1) 调优目标是比较特征提取器而非优化单一模型，逐模型调优会引入混淆变量；"
         "(2) 80 例数据下网格搜索拟合的是折间噪声而非真实信号；"
         "(3) 固定超参是 CONCH、UNI 等多项同类基准研究的通行做法，便于跨研究比较。"),
        ("Q：80 例样本太少，结论是否可靠？",
         "A：(1) UVM 极端罕见，80 例已是 TCGA 全量数据，无法通过纳入更多样本解决；"
         "(2) 5 折 CV 折间方差较大属预期现象，我们在图5 和图7 完整展示了稳定性，"
         "未做掩盖；(3) Q2 结论在此样本量下达到 p<0.001，说明效应量足够大，"
         "结论可信。"),
        ("Q：为什么不进行外部验证？",
         "A：UVM 全球可获取的有 SCNA 分型标注的公开队列仅 TCGA，"
         "外部队列缺乏统一的 D3/M3 分型注释，是客观限制而非方法选择，"
         "应在论文讨论部分明确声明。"),
        ("Q：Phikon/MUSK 的预训练数据分组是否准确？",
         "A：依据已发表论文的原始声明分组：Phikon 使用 TCGA（tcga_pretrained=True），"
         "MUSK 使用非 TCGA 私有数据（tcga_pretrained=False）。"
         "若评审存疑，可在表注中加入原始文献引用。"),
    ]
    for q_text, a_text in qas:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.left_indent  = Cm(0.5)
        rq = p.add_run(q_text + "\n")
        _set(rq, 11, bold=True, color=(68, 114, 196))
        ra = p.add_run(a_text)
        _set(ra, 11)

    # -----------------------------------------------------------------------
    h1(doc, "附录：图表推荐放置位置（论文结构参考）")
    placements = [
        ("Methods",        "图S1（流程图）、Table 2（模型概览）"),
        ("Results §1",     "Table 1（临床特征）、图1（临床分布）"),
        ("Results §2",     "Table 3（完整排名）、图2（Q1 柱状图）、图6（ROC 曲线）"),
        ("Results §3",     "图3（Q2 分组比较）、Table 5（Q2 三线表）"),
        ("Results §4",     "图4（Q3 交互图）、Table 6（Q3 三线表）"),
        ("Discussion",     "图10（范式分析）、图8（架构散点图）"),
        ("Supplementary",  "图5（箱线图）、图7（稳定性散点）、图9（指标相关）、"
                           "图11（多指标柱图）、Table 4（ΔAUC）"),
    ]
    for sec_name, figs in placements:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.left_indent  = Cm(0.5)
        rb = p.add_run(f"[{sec_name}]  ")
        _set(rb, 10.5, bold=True, color=(31, 73, 125))
        rv = p.add_run(figs)
        _set(rv, 10.5)

    doc.save(os.path.join(OUT_DIR, "thesis_narrative.docx"))
    print("  Saved: thesis_narrative.docx")


# ===========================================================================
if __name__ == "__main__":
    print(f"Output: {OUT_DIR}\n")
    print("[Doc 1] Bilingual figure captions ...")
    make_captions_doc()
    print("[Doc 2] Thesis narrative ...")
    make_narrative_doc()
    print("\nDone.")
